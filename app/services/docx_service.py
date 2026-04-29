import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# =========================
# 🧩 HELPER: NUMERACIÓN ISO
# =========================
def add_page_number(paragraph):
    run = paragraph.add_run()

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_total_pages(paragraph):
    run = paragraph.add_run()

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = "NUMPAGES"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# =========================
# 🧱 BORDES TABLA (FIX PRO)
# =========================
def set_table_borders(table):
    tbl = table._element
    tblPr = tbl.tblPr

    borders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)

    tblPr.append(borders)


# =========================
# 🏢 HEADER CORPORATIVO
# =========================
def configurar_header(doc, empresa):

    section = doc.sections[0]
    header = section.header

    table = header.add_table(rows=1, cols=2, width=Inches(6))

    cell_logo = table.rows[0].cells[0]
    try:
        cell_logo.paragraphs[0].add_run().add_picture(
            "app/assets/logo.png", width=Inches(1.2)
        )
    except:
        cell_logo.text = "LOGO"

    cell_text = table.rows[0].cells[1]
    p = cell_text.paragraphs[0]
    p.text = f"{empresa}\nSistema NOM-035"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


# =========================
# 📄 FOOTER ISO
# =========================
def configurar_footer(doc):

    section = doc.sections[0]
    footer = section.footer

    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p.add_run("Página ")
    add_page_number(p)
    p.add_run(" de ")
    add_total_pages(p)


# =========================
# 🎨 ESTILOS
# =========================
def configurar_estilos(doc):

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)


# =========================
# 🎯 RESPONSABLE DINÁMICO
# =========================
def obtener_responsable(dominio):
    d = dominio.lower()

    if "liderazgo" in d:
        return "Dirección / Gerencia"

    elif "relaciones" in d:
        return "Recursos Humanos"

    elif "ambiente" in d or "condiciones" in d:
        return "Seguridad e Higiene"

    elif "jornada" in d:
        return "Operaciones"

    elif "control" in d:
        return "Operaciones"

    elif "interferencia" in d or "familia" in d:
        return "Recursos Humanos"

    return "Recursos Humanos"


# =========================
# 🧾 POLÍTICAS PRO
# =========================
def generar_politicas_docx(path, empresa):

    doc = Document()

    configurar_estilos(doc)
    configurar_header(doc, empresa)
    configurar_footer(doc)

    # =========================
    # 📌 PORTADA
    # =========================
    titulo = doc.add_paragraph()
    run = titulo.add_run("POLÍTICA DE PREVENCIÓN DE RIESGOS PSICOSOCIALES")
    run.bold = True
    run.font.size = Pt(16)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    doc.add_paragraph(f"Empresa: {empresa}")
    doc.add_paragraph("Código: NOM035-POL-01")
    doc.add_paragraph("Versión: 1.0")
    doc.add_paragraph("Fecha de emisión: ____________________")
    doc.add_paragraph("Revisión: Anual")

    doc.add_paragraph("")

    # =========================
    # 1. OBJETIVO
    # =========================
    doc.add_heading('1. Objetivo', 1)
    doc.add_paragraph(
        "Establecer los lineamientos para identificar, analizar y prevenir los factores de riesgo psicosocial, "
        "así como para promover un entorno organizacional favorable, en cumplimiento con la NOM-035-STPS-2018."
    )

    # =========================
    # 2. ALCANCE
    # =========================
    doc.add_heading('2. Alcance', 1)
    doc.add_paragraph(
        "Aplica a todos los trabajadores del centro de trabajo, sin distinción de puesto, área o tipo de contratación."
    )

    # =========================
    # 3. MARCO NORMATIVO
    # =========================
    doc.add_heading('3. Marco Normativo', 1)
    doc.add_paragraph("• NOM-035-STPS-2018")
    doc.add_paragraph("• Ley Federal del Trabajo")
    doc.add_paragraph("• Reglamento Federal de Seguridad y Salud en el Trabajo")

    # =========================
    # 4. DEFINICIONES
    # =========================
    doc.add_heading('4. Definiciones', 1)
    doc.add_paragraph("Factores de riesgo psicosocial: Aquellos que pueden provocar trastornos de ansiedad, estrés o alteraciones del ciclo sueño-vigilia.")
    doc.add_paragraph("Entorno organizacional favorable: Aquel en el que se promueve el sentido de pertenencia, formación adecuada y cargas equilibradas.")
    doc.add_paragraph("Violencia laboral: Actos de hostigamiento, acoso o malos tratos dentro del entorno laboral.")

    # =========================
    # 5. RESPONSABILIDADES
    # =========================
    doc.add_heading('5. Responsabilidades', 1)

    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Rol"
    table.rows[0].cells[1].text = "Responsabilidad"

    roles = [
        ("Dirección General", "Aprobar y garantizar el cumplimiento de la NOM-035"),
        ("Recursos Humanos", "Implementar evaluaciones, medidas preventivas y seguimiento"),
        ("Mandos medios", "Aplicar acciones y supervisar condiciones laborales"),
        ("Trabajadores", "Participar en evaluaciones y reportar situaciones de riesgo")
    ]

    for rol, resp in roles:
        row = table.add_row().cells
        row[0].text = rol
        row[1].text = resp

    # =========================
    # 6. PROCEDIMIENTO
    # =========================
    doc.add_heading('6. Procedimiento', 1)

    doc.add_paragraph("6.1 Identificación de riesgos")
    doc.add_paragraph("• Aplicación de cuestionarios NOM-035")
    doc.add_paragraph("• Análisis de resultados")

    doc.add_paragraph("6.2 Evaluación")
    doc.add_paragraph("• Clasificación de niveles de riesgo")
    doc.add_paragraph("• Identificación de dominios críticos")

    doc.add_paragraph("6.3 Implementación de acciones")
    doc.add_paragraph("• Plan de acción correctivo y preventivo")
    doc.add_paragraph("• Capacitación y ajustes organizacionales")

    doc.add_paragraph("6.4 Seguimiento")
    doc.add_paragraph("• Monitoreo continuo")
    doc.add_paragraph("• Reevaluación periódica")

    # =========================
    # 7. CANAL DE QUEJAS
    # =========================
    doc.add_heading('7. Canal de Quejas', 1)
    doc.add_paragraph(
        "La empresa cuenta con un mecanismo confidencial para la recepción de quejas relacionadas con factores de riesgo psicosocial y violencia laboral."
    )
    doc.add_paragraph("Medios disponibles:")
    doc.add_paragraph("• Correo electrónico: ____________________")
    doc.add_paragraph("• Buzón físico")
    doc.add_paragraph("• Recursos Humanos")

    # =========================
    # 8. MEDIDAS PREVENTIVAS
    # =========================
    doc.add_heading('8. Medidas Preventivas', 1)
    doc.add_paragraph("• Promoción del equilibrio vida-trabajo")
    doc.add_paragraph("• Capacitación en liderazgo")
    doc.add_paragraph("• Prevención de violencia laboral")
    doc.add_paragraph("• Mejora de condiciones organizacionales")

    # =========================
    # 9. CONTROL DOCUMENTAL
    # =========================
    doc.add_heading('9. Control Documental', 1)
    doc.add_paragraph("Este documento deberá ser revisado al menos una vez al año o cuando existan cambios en la normativa aplicable.")

    doc.add_paragraph("")
    doc.add_paragraph("__________________________")
    doc.add_paragraph("Firma Dirección General")

    doc.save(path)


# =========================
# 🛠️ PLAN DE ACCIÓN PRO (FIX)
# =========================
def generar_plan_docx(path, plan_acciones, empresa):

    doc = Document()

    configurar_estilos(doc)
    configurar_header(doc, empresa)
    configurar_footer(doc)

    doc.add_heading("PLAN DE ACCIÓN NOM-035", 0)

    doc.add_paragraph(f"Empresa: {empresa}")
    doc.add_paragraph("Fecha de emisión: ____________________")
    doc.add_paragraph("Versión: 1.0")

    doc.add_paragraph("")
    doc.add_paragraph(
        "El presente plan de acción tiene como objetivo establecer las medidas "
        "preventivas y correctivas derivadas de la evaluación de factores de riesgo "
        "psicosocial conforme a la NOM-035-STPS-2018, priorizando los dominios con mayor nivel de riesgo."
    )

    # =========================
    # 📊 TABLA RESUMEN
    # =========================
    doc.add_heading("1. Resumen de Acciones", 1)

    table = doc.add_table(rows=1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = [
        "Dominio", "Nivel", "Prioridad", "Acción",
        "Responsable", "Inicio", "Fin", "Evidencia"
    ]

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Fondo oscuro
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), "1F2937")
        tcPr.append(shd)

    # =========================
    # 🔥 PRIORIDAD AUTOMÁTICA
    # =========================
    def calcular_prioridad(nivel):
        if nivel in ["Alto", "Muy alto"]:
            return "Alta"
        elif nivel == "Medio":
            return "Media"
        else:
            return "Baja"

    for p in plan_acciones:
        row = table.add_row().cells

        prioridad = calcular_prioridad(p["nivel"])

        # ✅ FIX RESPONSABLE
        responsable = obtener_responsable(p["dominio"])

        row[0].text = p["dominio"]
        row[1].text = p["nivel"]
        row[2].text = prioridad
        row[3].text = p["accion"]
        row[4].text = responsable
        row[5].text = p.get("inicio", "Mes 1")
        row[6].text = p.get("fin", "Mes 3")
        row[7].text = p.get(
            "evidencia",
            "Actas / KPIs / Reportes / Evidencia fotográfica"
        )

    # ✅ FIX BORDES
    set_table_borders(table)

    doc.add_paragraph("")

    # =========================
    # 🔍 DETALLE
    # =========================
    doc.add_heading("2. Detalle de Acciones", 1)

    for p in plan_acciones:

        prioridad = calcular_prioridad(p["nivel"])
        responsable = obtener_responsable(p["dominio"])

        doc.add_heading(p["dominio"], 2)

        doc.add_paragraph(f"Nivel de riesgo: {p['nivel']}")
        doc.add_paragraph(f"Prioridad: {prioridad}")

        doc.add_paragraph("Acción a implementar:")
        doc.add_paragraph(p["accion"], style="List Bullet")

        doc.add_paragraph("Responsable:")
        doc.add_paragraph(responsable, style="List Bullet")

        doc.add_paragraph("Plazo:")
        doc.add_paragraph(p.get("plazo", "30-90 días"), style="List Bullet")

        doc.add_paragraph("Indicador (KPI):")
        doc.add_paragraph(
            "Reducción del puntaje en evaluación posterior ≥ 10%",
            style="List Bullet"
        )

        doc.add_paragraph("Evidencia requerida:")
        doc.add_paragraph(
            "Actas firmadas, reportes mensuales, indicadores de seguimiento y evidencia documental",
            style="List Bullet"
        )

        doc.add_paragraph("")

    # =========================
    # ⚠️ CUMPLIMIENTO
    # =========================
    doc.add_heading("3. Cumplimiento NOM-035", 1)

    obligaciones = [
        "Implementar acciones de control sobre los factores de riesgo identificados",
        "Registrar evidencia documental de las acciones realizadas",
        "Difundir políticas de prevención de riesgos psicosociales",
        "Realizar evaluaciones periódicas",
        "Dar seguimiento a indicadores establecidos"
    ]

    for o in obligaciones:
        doc.add_paragraph(o, style="List Bullet")

    # =========================
    # 🎯 CIERRE
    # =========================
    doc.add_heading("4. Conclusión", 1)

    doc.add_paragraph(
        "La correcta implementación del presente plan permitirá reducir los factores de riesgo psicosocial, "
        "mejorar el entorno organizacional y asegurar el cumplimiento de la NOM-035-STPS-2018."
    )

    doc.add_paragraph("")

    doc.add_heading("Seguimiento y Control", 1)
    doc.add_paragraph(
        "El área de Recursos Humanos deberá dar seguimiento periódico al cumplimiento "
        "de las acciones establecidas en este plan."
    )

    doc.save(path)

# =========================
# 📁 GENERAR ENTREGABLES NOM-035
# =========================
def generar_entregables_nom035(base_path, empresa, interpretacion):

    os.makedirs(base_path, exist_ok=True)

    archivos_generados = []

    # =========================
    # 🧾 1. ACTA DE IMPLEMENTACIÓN
    # =========================
    path_acta = os.path.join(base_path, "01_Acta_Implementacion.docx")
    generar_acta_docx(path_acta, empresa)
    archivos_generados.append(path_acta)

    # =========================
    # 📊 2. SEGUIMIENTO KPI
    # =========================
    path_kpi = os.path.join(base_path, "02_Seguimiento_KPI.docx")
    generar_kpi_docx(path_kpi, empresa)
    archivos_generados.append(path_kpi)

    # =========================
    # ⚠️ 3. REPORTE DE INCIDENTES
    # =========================
    path_incidentes = os.path.join(base_path, "03_Reporte_Incidentes.docx")
    generar_incidentes_docx(path_incidentes, empresa)
    archivos_generados.append(path_incidentes)

    # =========================
    # 🧠 4. MINUTA DE CAPACITACIÓN
    # =========================
    path_minuta = os.path.join(base_path, "04_Minuta_Capacitacion.docx")
    generar_minuta_docx(path_minuta, empresa)
    archivos_generados.append(path_minuta)

    return archivos_generados


def generar_acta_docx(path, empresa):

    doc = Document()
    configurar_estilos(doc)
    configurar_header(doc, empresa)
    configurar_footer(doc)

    doc.add_heading("ACTA DE IMPLEMENTACIÓN NOM-035", 0)

    doc.add_paragraph(f"Empresa: {empresa}")
    doc.add_paragraph("Fecha: ____________________")
    doc.add_paragraph("Lugar: ____________________")

    doc.add_paragraph("")
    doc.add_heading("Instrucciones de uso", 1)
    doc.add_paragraph(
        "Este documento formaliza el inicio del plan de acción NOM-035. "
        "Debe ser firmado por Dirección y Recursos Humanos como evidencia de cumplimiento."
    )

    doc.add_heading("1. Objetivo", 1)
    doc.add_paragraph(
        "Establecer formalmente el inicio de acciones para la prevención de riesgos psicosociales."
    )

    doc.add_heading("2. Alcance", 1)
    doc.add_paragraph("Aplica a todas las áreas evaluadas.")

    doc.add_heading("3. Acuerdos establecidos", 1)
    acuerdos = [
        "Aprobación del plan de acción",
        "Asignación de responsables",
        "Seguimiento mensual"
    ]
    for a in acuerdos:
        doc.add_paragraph(a, style="List Bullet")

    doc.add_heading("4. Responsables", 1)
    table = doc.add_table(rows=1, cols=3)

    headers = ["Área", "Responsable", "Firma"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for area in ["Dirección", "RH", "Operaciones"]:
        row = table.add_row().cells
        row[0].text = area
        row[1].text = "________________"
        row[2].text = ""

    doc.add_paragraph("")
    doc.add_heading("5. Observaciones", 1)
    doc.add_paragraph("________________________________________")

    doc.save(path)


def generar_kpi_docx(path, empresa):

    doc = Document()
    configurar_estilos(doc)
    configurar_header(doc, empresa)
    configurar_footer(doc)

    doc.add_heading("SEGUIMIENTO DE INDICADORES NOM-035", 0)

    doc.add_paragraph(f"Empresa: {empresa}")

    doc.add_heading("Instrucciones de uso", 1)
    doc.add_paragraph(
        "Actualizar este formato mensualmente. "
        "Registrar avances reales y anexar evidencia documental."
    )

    doc.add_heading("Periodo evaluado", 1)
    doc.add_paragraph("Mes: ____________________")

    table = doc.add_table(rows=1, cols=6)

    headers = ["Indicador", "Meta", "Resultado", "Avance %", "Estatus", "Evidencia"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        run = cell.paragraphs[0].add_run(h)
        run.bold = True

    indicadores = [
        "Reducción de riesgo",
        "Clima laboral",
        "Rotación",
        "Ausentismo",
        "Cumplimiento del plan"
    ]

    for ind in indicadores:
        row = table.add_row().cells
        row[0].text = ind
        row[1].text = "Definir"
        row[2].text = ""
        row[3].text = ""
        row[4].text = "Pendiente"
        row[5].text = ""

    doc.add_paragraph("")
    doc.add_heading("Conclusión del periodo", 1)
    doc.add_paragraph("________________________________________")

    doc.save(path)


def generar_incidentes_docx(path, empresa):

    doc = Document()
    configurar_estilos(doc)
    configurar_header(doc, empresa)
    configurar_footer(doc)

    doc.add_heading("REPORTE DE INCIDENTES PSICOSOCIALES", 0)

    doc.add_paragraph(f"Empresa: {empresa}")

    doc.add_heading("Instrucciones", 1)
    doc.add_paragraph(
        "Registrar cualquier evento relacionado con estrés, carga laboral o violencia. "
        "Este documento sirve como evidencia ante auditoría."
    )

    doc.add_heading("1. Datos generales", 1)
    doc.add_paragraph("Fecha: ____________________")
    doc.add_paragraph("Área: ____________________")
    doc.add_paragraph("Reporta: ____________________")

    doc.add_heading("2. Descripción", 1)
    doc.add_paragraph("________________________________________")

    doc.add_heading("3. Nivel de riesgo percibido", 1)
    doc.add_paragraph("Bajo / Medio / Alto")

    doc.add_heading("4. Acciones tomadas", 1)
    doc.add_paragraph("________________________________________")

    doc.add_heading("5. Seguimiento", 1)
    doc.add_paragraph("________________________________________")

    doc.add_heading("6. Responsable", 1)
    doc.add_paragraph("____________________")

    doc.save(path)


def generar_minuta_docx(path, empresa):

    doc = Document()
    configurar_estilos(doc)
    configurar_header(doc, empresa)
    configurar_footer(doc)

    doc.add_heading("MINUTA DE CAPACITACIÓN NOM-035", 0)

    doc.add_paragraph(f"Empresa: {empresa}")

    doc.add_heading("Instrucciones", 1)
    doc.add_paragraph(
        "Utilizar este formato para documentar capacitaciones relacionadas con NOM-035. "
        "Debe incluir lista de asistencia y conclusiones."
    )

    doc.add_heading("1. Datos generales", 1)
    doc.add_paragraph("Fecha: ____________________")
    doc.add_paragraph("Instructor: ____________________")
    doc.add_paragraph("Duración: ____________________")

    doc.add_heading("2. Tema", 1)
    doc.add_paragraph("________________________________________")

    doc.add_heading("3. Objetivo", 1)
    doc.add_paragraph("________________________________________")

    doc.add_heading("4. Lista de asistencia", 1)

    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "Nombre"
    table.rows[0].cells[1].text = "Área"
    table.rows[0].cells[2].text = "Firma"

    for _ in range(6):
        row = table.add_row().cells
        row[0].text = ""
        row[1].text = ""
        row[2].text = ""

    doc.add_heading("5. Conclusiones", 1)
    doc.add_paragraph("________________________________________")

    doc.add_heading("6. Compromisos", 1)
    doc.add_paragraph("• ____________________")

    doc.save(path)